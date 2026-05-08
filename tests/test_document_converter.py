"""
Tests for document_converter.py
"""

import pytest
from services.document_converter import convert_document


class TestConvertText:
    def test_plain_text_wrapped_in_pre(self):
        doc = convert_document(b"Hello World", "test.txt")
        assert doc.html == "<pre class=\"doc-content\">Hello World</pre>"
        assert doc.text == "Hello World"
        assert doc.content_type == "text/plain"

    def test_big5_fallback(self):
        raw = "中文測試".encode("big5")
        doc = convert_document(raw, "test.txt")
        assert "中文測試" in doc.html

    def test_unknown_extension_treated_as_text(self):
        doc = convert_document(b"some data", "file.bin")
        assert doc.html.startswith("<pre")

    def test_no_extension_treated_as_text(self):
        doc = convert_document(b"no extension", "README")
        assert doc.html.startswith("<pre")

    def test_html_escaped(self):
        doc = convert_document(b"<script>alert('xss')</script>", "test.txt")
        assert "&lt;script&gt;" in doc.html
        assert "<script>" not in doc.html


class TestConvertMarkdown:
    def test_basic_markdown(self):
        raw = b"# Hello\n\nThis is **bold** text."
        doc = convert_document(raw, "test.md")
        assert "<h1>" in doc.html
        assert "<strong>" in doc.html
        assert doc.content_type == "text/markdown"

    def test_code_block(self):
        raw = b"```python\nprint('hello')\n```"
        doc = convert_document(raw, "test.md")
        assert "<code>" in doc.html or "<pre>" in doc.html


class TestConvertDocx:
    def test_simple_docx(self):
        from docx import Document
        from io import BytesIO
        docx = Document()
        docx.add_paragraph("Hello World")
        docx.add_paragraph("Second paragraph")
        buf = BytesIO()
        docx.save(buf)
        raw = buf.getvalue()

        doc = convert_document(raw, "test.docx")
        assert "Hello World" in doc.html
        assert "Second paragraph" in doc.html

    def test_docx_has_paragraphs(self):
        from docx import Document
        from io import BytesIO
        docx = Document()
        docx.add_paragraph("First")
        docx.add_paragraph("Second")
        buf = BytesIO()
        docx.save(buf)

        doc = convert_document(buf.getvalue(), "test.docx")
        # mammoth wraps paragraphs in <p>
        assert "<p>" in doc.html


class TestConvertPdf:
    def test_simple_pdf(self):
        from reportlab.pdfgen import canvas
        from io import BytesIO
        buf = BytesIO()
        c = canvas.Canvas(buf)
        c.drawString(100, 750, "Hello PDF World")
        c.save()
        raw = buf.getvalue()

        doc = convert_document(raw, "test.pdf")
        assert "Hello PDF World" in doc.html or "Hello" in doc.html
        assert "<p>" in doc.html or "<div>" in doc.html
